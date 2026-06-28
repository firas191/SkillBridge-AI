"""Document parsing for real CV / job files.

Supported inputs (all read fully offline, no API key):
  * PDF   — digital text via pdfplumber; falls back to OCR for scanned PDFs.
  * DOCX  — python-docx (paragraphs + tables).
  * TXT / MD — decoded directly.
  * Images (PNG/JPG/JPEG/WEBP/BMP/TIFF) — OCR.

OCR uses RapidOCR (PaddleOCR's detection/recognition models via ONNX Runtime):
accurate and robust to noisy real-world scans, with a pure-pip install and no
native binary. A light preprocessing pass (grayscale, auto-contrast, upscaling)
improves recognition on low-quality or photographed documents.
"""
from __future__ import annotations

import io
from dataclasses import dataclass
from pathlib import Path
from threading import Lock

from app.core.logging import get_logger

log = get_logger(__name__)

MAX_CHARS = 30_000  # guardrail to keep prompts within context limits
# If a PDF yields less text than this, treat it as scanned and OCR it.
_PDF_OCR_TEXT_THRESHOLD = 100
# Render scanned PDF pages at this DPI before OCR (higher = better, slower).
_PDF_OCR_DPI = 300

IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}
TEXT_SUFFIXES = {".txt", ".md", ".text"}


class OCRUnavailableError(RuntimeError):
    """Raised when an image/scanned doc needs OCR but the engine isn't installed."""


@dataclass
class ParsedDocument:
    text: str
    filename: str
    char_count: int
    method: str  # pdf-text | pdf-ocr | docx | text | image-ocr


# ---------------------------------------------------------------------------
# OCR engine (lazy singleton — model load is expensive, so reuse it)
# ---------------------------------------------------------------------------
_ocr_engine = None
_ocr_lock = Lock()


def _get_ocr():
    global _ocr_engine
    if _ocr_engine is None:
        with _ocr_lock:
            if _ocr_engine is None:
                try:
                    from rapidocr_onnxruntime import RapidOCR
                except ImportError as exc:  # pragma: no cover - env dependent
                    raise OCRUnavailableError(
                        "OCR engine not installed. Run "
                        "`pip install -r requirements.txt` (rapidocr-onnxruntime)."
                    ) from exc
                log.info("Initialising RapidOCR engine (first use)...")
                _ocr_engine = RapidOCR()
    return _ocr_engine


def _preprocess(img):
    """Light, accuracy-preserving cleanup for noisy real-world scans."""
    from PIL import Image, ImageOps

    img = img.convert("RGB")
    # Upscale small/low-DPI images so thin strokes survive detection.
    longest = max(img.size)
    if longest < 1600:
        scale = 1600 / longest
        new_size = (int(img.size[0] * scale), int(img.size[1] * scale))
        img = img.resize(new_size, Image.LANCZOS)
    # Auto-contrast helps faint photocopies/photos without destroying color.
    gray = ImageOps.grayscale(img)
    gray = ImageOps.autocontrast(gray, cutoff=1)
    return gray.convert("RGB")


def _ocr_pil_image(img) -> str:
    import numpy as np

    engine = _get_ocr()
    processed = _preprocess(img)
    result, _ = engine(np.asarray(processed))
    if not result:
        return ""
    # RapidOCR returns [[box, text, score], ...] in reading order.
    return "\n".join(line[1] for line in result if len(line) >= 2 and line[1])


def _ocr_image_bytes(data: bytes) -> str:
    from PIL import Image

    with Image.open(io.BytesIO(data)) as img:
        return _ocr_pil_image(img)


# ---------------------------------------------------------------------------
# Per-format extractors
# ---------------------------------------------------------------------------
def _from_pdf(data: bytes) -> tuple[str, str]:
    """Return (text, method). OCR fallback for scanned PDFs."""
    import pdfplumber

    pages: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            pages.append(page.extract_text() or "")
    text = "\n".join(pages).strip()

    if len(text) >= _PDF_OCR_TEXT_THRESHOLD:
        return text, "pdf-text"

    # Likely a scanned/image PDF — render each page and OCR it.
    log.info("PDF has little extractable text (%d chars); using OCR fallback.", len(text))
    ocr_text = _ocr_pdf_pages(data)
    if ocr_text.strip():
        return ocr_text.strip(), "pdf-ocr"
    # Return whatever digital text we had (may be empty) rather than failing hard.
    return text, "pdf-text"


def _ocr_pdf_pages(data: bytes) -> str:
    import fitz  # PyMuPDF
    from PIL import Image

    chunks: list[str] = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        for page in doc:
            pix = page.get_pixmap(dpi=_PDF_OCR_DPI, alpha=False)
            img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
            chunks.append(_ocr_pil_image(img))
    return "\n".join(chunks)


def _from_docx(data: bytes) -> str:
    import docx  # python-docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:  # skills are often in tables
        for row in table.rows:
            parts.append(" \t ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Public entrypoint
# ---------------------------------------------------------------------------
def extract_text(data: bytes, filename: str) -> ParsedDocument:
    """Extract plain text from uploaded document bytes.

    Raises ``ValueError`` for unsupported types and ``OCRUnavailableError`` when
    OCR is required but the engine isn't installed.
    """
    suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        text, method = _from_pdf(data)
    elif suffix == ".docx":
        text, method = _from_docx(data), "docx"
    elif suffix in TEXT_SUFFIXES:
        text, method = data.decode("utf-8", errors="replace"), "text"
    elif suffix in IMAGE_SUFFIXES:
        text, method = _ocr_image_bytes(data), "image-ocr"
    else:
        raise ValueError(
            f"Unsupported file type '{suffix}'. "
            "Use PDF, DOCX, TXT, MD, or an image (PNG/JPG/WEBP/BMP/TIFF)."
        )

    text = (text or "").strip()
    if len(text) > MAX_CHARS:
        log.info("Truncating %s from %d to %d chars", filename, len(text), MAX_CHARS)
        text = text[:MAX_CHARS]
    return ParsedDocument(
        text=text, filename=filename, char_count=len(text), method=method
    )
